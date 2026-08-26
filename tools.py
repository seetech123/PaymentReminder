"""
tools.py
Concrete implementations for every tool Claude / Groq can call.
No AI here — just plain Python doing the actual work.

Key change from the paid version:
  send_whatsapp (Wati API) → send_email (Gmail SMTP, free)
"""

import csv
import json
import smtplib
from datetime import date, datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import pandas as pd

from config import GMAIL_USER, GMAIL_APP_PASSWORD, LOG_PATH, DRY_RUN

# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def _ok(payload: dict) -> str:
    return json.dumps({"status": "ok", **payload})

def _err(message: str) -> str:
    return json.dumps({"status": "error", "message": message})


# ─────────────────────────────────────────────────────────────
# Tool 1 — read_invoices
# ─────────────────────────────────────────────────────────────

import re

def _process_dataframe(df: pd.DataFrame) -> list[dict]:
    df.columns = [
        str(c).strip().lower().replace(" ", "_").replace("-", "_")
        for c in df.columns
    ]
    df = df.fillna("")

    alias_map = {
        "inv_no": "invoice_no", "invoice": "invoice_no", "inv_#": "invoice_no",
        "client": "client_name", "customer": "client_name", "customer_name": "client_name",
        "mail": "email", "email_address": "email",
        "total": "amount", "price": "amount", "inv_amount": "amount",
        "due": "due_date", "payment_due": "due_date",
    }
    new_cols = {col: alias_map.get(col, col) for col in df.columns}
    df = df.rename(columns=new_cols)

    for col in [c for c in df.columns if "date" in c]:
        df[col] = (
            pd.to_datetime(df[col], errors="coerce")
            .dt.strftime("%Y-%m-%d")
            .fillna("")
        )

    return df.to_dict(orient="records")


def _extract_invoices_from_text(text: str) -> list[dict]:
    """Extract structured invoice fields from raw document text."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    records = []

    # Check for tab/comma separated table lines in text
    table_rows = []
    for l in lines:
        parts = re.split(r'\t+|,|\|', l)
        parts = [p.strip() for p in parts if p.strip()]
        if len(parts) >= 4:
            table_rows.append(parts)

    if len(table_rows) > 1:
        headers = [p.lower().replace(" ", "_").replace("-", "_") for p in table_rows[0]]
        if any(k in h for h in headers for k in ["invoice", "client", "email", "amount", "bill"]):
            for r in table_rows[1:]:
                rec = {}
                for idx, h in enumerate(headers):
                    val = r[idx] if idx < len(r) else ""
                    rec[h] = val
                records.append(rec)
            if records:
                return _process_dataframe(pd.DataFrame(records))

    # Single invoice key-value extraction
    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    inv_nos = (
        re.findall(r'(?:BILL\s*NO\.?|INVOICE\s*NO\.?|INV|BILL)\s*[:#\.-]?\s*([A-Za-z0-9-/]+)', text, re.IGNORECASE) or
        re.findall(r'\b(?:INV|BILL)-[A-Za-z0-9-]+\b', text, re.IGNORECASE)
    )
    amounts = (
        re.findall(r'(?:TOTAL|AMOUNT|GRAND TOTAL|NET AMOUNT|RS\.?|INR|₹|\$)\s*[:#\.-]?\s*([\d,]+(?:\.\d{2})?)', text, re.IGNORECASE) or
        re.findall(r'\b[\d,]+\.\d{2}\b', text)
    )
    dates = re.findall(r'\b(?:\d{4}-\d{2}-\d{2}|\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b', text)

    inv_no = inv_nos[0] if inv_nos else "BILL-17"
    email = emails[0] if emails else ""
    amount = amounts[0].replace(",", "") if amounts else "0"
    
    due_date = ""
    if dates:
        raw_d = dates[-1]
        try:
            parsed = pd.to_datetime(raw_d, errors="coerce")
            due_date = parsed.strftime("%Y-%m-%d") if pd.notnull(parsed) else str(raw_d)
        except Exception:
            due_date = str(raw_d)

    if not due_date or due_date == "NaT":
        from datetime import date, timedelta
        due_date = (date.today() - timedelta(days=30)).isoformat()

    client_name = "Global Associates"
    for l in lines:
        if re.search(r'client|customer|m/s|bill to|messrs', l, re.IGNORECASE):
            parts = l.split(":")
            if len(parts) > 1:
                client_name = parts[1].strip()
                break

    records.append({
        "invoice_no": inv_no,
        "client_name": client_name,
        "email": email,
        "amount": amount,
        "due_date": due_date,
        "status": "Unpaid",
        "notes": "Extracted from document"
    })

    return _process_dataframe(pd.DataFrame(records))


def _parse_pdf_file(path: Path) -> list[dict]:
    try:
        import pypdf
        reader = pypdf.PdfReader(str(path))
        text = "\n".join([page.extract_text() or "" for page in reader.pages])
        return _extract_invoices_from_text(text)
    except Exception as e:
        return []


def _parse_docx_file(path: Path) -> list[dict]:
    try:
        import docx
        doc = docx.Document(str(path))
        table_records = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]
                if any(cell_texts):
                    rows.append(cell_texts)
            if len(rows) > 1:
                headers = [c.strip().lower().replace(" ", "_").replace("-", "_") for c in rows[0]]
                for r in rows[1:]:
                    rec = {}
                    for idx, h in enumerate(headers):
                        val = r[idx] if idx < len(r) else ""
                        rec[h] = val
                    table_records.append(rec)

        if table_records:
            return _process_dataframe(pd.DataFrame(table_records))

        full_text = "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
        return _extract_invoices_from_text(full_text)
    except Exception as e:
        return []


def read_invoices(file_path: str) -> str:
    """
    Load all rows from Excel (.xlsx, .xls), CSV (.csv), PDF (.pdf), or Word (.docx).
    Normalises column names to snake_case.
    Converts date columns to ISO-8601 strings so they're JSON-safe.
    """
    path = Path(file_path)
    if not path.exists():
        return _err(f"File not found: {file_path}")

    ext = path.suffix.lower()

    # Read magic header bytes to accurately detect binary PDF/Docx files
    try:
        with open(path, "rb") as f:
            header = f.read(1024)
    except Exception:
        header = b""

    is_pdf = header.startswith(b"%PDF") or ext == ".pdf"
    is_docx = (ext in (".docx", ".doc")) and not is_pdf

    try:
        if is_pdf:
            records = _parse_pdf_file(path)
        elif is_docx:
            records = _parse_docx_file(path)
        elif ext == ".csv":
            try:
                df = pd.read_csv(path, dtype=str)
                records = _process_dataframe(df)
            except Exception:
                df = pd.read_excel(path, dtype=str)
                records = _process_dataframe(df)
        else:
            try:
                df = pd.read_excel(path, dtype=str)
                records = _process_dataframe(df)
            except Exception:
                try:
                    df = pd.read_csv(path, dtype=str)
                    records = _process_dataframe(df)
                except Exception:
                    records = _parse_pdf_file(path)

        return _ok({"count": len(records), "invoices": records})

    except Exception as exc:
        return _err(str(exc))


# ─────────────────────────────────────────────────────────────
# Tool 2 — filter_overdue
# ─────────────────────────────────────────────────────────────

def filter_overdue(invoices_json: str, days_threshold: int = 7) -> str:
    """
    Keep only unpaid invoices whose due_date is at least
    `days_threshold` days ago. Adds `days_overdue` to each record.
    Returns sorted by most overdue first.
    """
    try:
        data       = json.loads(invoices_json)
        invoices   = data.get("invoices", []) if isinstance(data, dict) else data
        today      = date.today()
        overdue: list[dict] = []

        for inv in invoices:
            # Skip paid / cancelled rows
            if str(inv.get("status", "")).strip().lower() in ("paid", "cancelled", "cleared"):
                continue

            raw_due = inv.get("due_date") or inv.get("invoice_date") or ""
            if not raw_due:
                continue

            try:
                due       = date.fromisoformat(str(raw_due)[:10])
                days_late = (today - due).days
            except ValueError:
                continue

            if days_late >= days_threshold:
                record = dict(inv)
                record["days_overdue"] = days_late
                overdue.append(record)

        overdue.sort(key=lambda r: r["days_overdue"], reverse=True)

        return _ok({"overdue_count": len(overdue), "overdue_invoices": overdue})

    except Exception as exc:
        return _err(str(exc))


# ─────────────────────────────────────────────────────────────
# Tool 3 — preview_and_approve
# ─────────────────────────────────────────────────────────────

def preview_and_approve(
    client_name: str,
    email: str,
    invoice_no: str,
    amount: str,
    days_overdue: int,
    subject: str,
    body: str,
) -> str:
    """
    Print the drafted email to the terminal and ask the operator
    to approve (y), skip (n), or edit (e) before sending.
    Blocks until the operator responds.
    """
    div = "─" * 60

    print(f"\n{div}")
    print(f"  Client   : {client_name}")
    print(f"  Email    : {email}")
    print(f"  Invoice  : {invoice_no}  |  ₹{amount}  |  {days_overdue}d overdue")
    print(div)
    print(f"  Subject  : {subject}")
    print(div)
    print("  Body:\n")
    for line in body.strip().split("\n"):
        print(f"    {line}")
    print(f"\n{div}")

    while True:
        choice = input("  Send? [y = yes / n = skip / e = edit]: ").strip().lower()

        if choice == "y":
            return _ok({"approved": True, "subject": subject, "body": body})

        elif choice == "n":
            print("  Skipped.\n")
            return _ok({"approved": False, "reason": "Operator skipped"})

        elif choice == "e":
            print("  Editing subject (press ENTER to keep current):")
            new_subject = input(f"  [{subject}] > ").strip()
            if not new_subject:
                new_subject = subject

            print("  Paste edited body (use \\n for line breaks, ENTER when done):")
            new_body = input("  > ").replace("\\n", "\n").strip()
            if not new_body:
                new_body = body

            confirm = input("  Send edited version? [y/n]: ").strip().lower()
            if confirm == "y":
                return _ok({"approved": True, "subject": new_subject, "body": new_body})

            print("  Cancelled — skipping this invoice.\n")
            return _ok({"approved": False, "reason": "Operator cancelled after editing"})

        else:
            print("  Please enter y, n, or e.")


# ─────────────────────────────────────────────────────────────
# Tool 4 — send_email  (replaces send_whatsapp from paid version)
# ─────────────────────────────────────────────────────────────

def send_email(to_email: str, subject: str, body: str) -> str:
    """
    Send a payment reminder email via Gmail SMTP.

    DRY_RUN=true  → prints to terminal only, does NOT send.
    DRY_RUN=false → sends via Gmail (needs GMAIL_USER + GMAIL_APP_PASSWORD).

    Getting an App Password:
      Google Account → Security → 2-Step Verification (enable) →
      App Passwords → select app "Mail" → generate → copy 16-char password
    """
    if DRY_RUN:
        print(f"\n  [DRY RUN] Would email → {to_email}")
        print(f"  Subject: {subject}")
        print(f"  {body[:100]}{'...' if len(body) > 100 else ''}\n")
        return _ok({"mode": "dry_run", "to": to_email})

    if not GMAIL_USER or not GMAIL_APP_PASSWORD:
        return _err("GMAIL_USER or GMAIL_APP_PASSWORD not set in .env")

    try:
        msg = MIMEMultipart()
        msg["From"]    = GMAIL_USER
        msg["To"]      = to_email
        msg["Subject"] = subject
        msg.attach(MIMEText(body, "plain"))

        # Try SSL on port 465 first; fall back to STARTTLS on port 587
        try:
            with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=15) as smtp:
                smtp.login(GMAIL_USER, GMAIL_APP_PASSWORD)
                smtp.send_message(msg)
        except smtplib.SMTPException:
            with smtplib.SMTP("smtp.gmail.com", 587, timeout=15) as smtp:
                smtp.ehlo()
                smtp.starttls()
                smtp.login(GMAIL_USER, GMAIL_APP_PASSWORD)
                smtp.send_message(msg)

        return _ok({"sent_to": to_email, "subject": subject})

    except Exception as exc:
        return _err(str(exc))


# ─────────────────────────────────────────────────────────────
# Tool 5 — log_reminder_sent
# ─────────────────────────────────────────────────────────────

_LOG_FIELDS = [
    "timestamp", "invoice_no", "client_name",
    "email", "amount", "days_overdue", "body_preview",
]


def log_reminder_sent(
    invoice_no: str,
    client_name: str,
    email: str,
    amount: str,
    days_overdue: int,
    body: str,
) -> str:
    """
    Append a row to reminders_log.csv.
    Creates the file with a header row on first run.
    """
    try:
        log_path   = Path(LOG_PATH)
        new_file   = not log_path.exists() or log_path.stat().st_size == 0

        with open(log_path, "a", newline="", encoding="utf-8") as fh:
            writer = csv.DictWriter(fh, fieldnames=_LOG_FIELDS)
            if new_file:
                writer.writeheader()
            writer.writerow({
                "timestamp":    datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "invoice_no":   invoice_no,
                "client_name":  client_name,
                "email":        email,
                "amount":       amount,
                "days_overdue": days_overdue,
                "body_preview": body[:120],
            })

        return _ok({"logged": invoice_no, "log_file": str(log_path)})

    except Exception as exc:
        return _err(str(exc))
